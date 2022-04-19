from docx import Document
from docx.shared import Inches
file = open("Book1.csv", mode="r", encoding="utf-8-sig")
header = file.readline()
document = Document()
row = file.readline()
while row != "":
    
    
    row_list = row.split(",")
    name = row_list[1]

    document.add_heading('giấy mời họp phụ huynh',0)

    p = document.add_paragraph('học sinh:')
    p.add_run(name).bold=True
    p.add_run(" -lớp 10").bold=True
    p=document.add_paragraph('thời gian: ')
    p.add_run("9AM ngày 1/1/2022").bold=True
    p=document.add_paragraph('địa điểm: ')
    p.add_run("phòng 10A1 trường THPT Nông Cống II").bold=True
    document.add_heading('nội dung',level=1)
    document.add_paragraph('tổng kết năm học 2022 vừa qua và nhận xét kết quả của từng học sinh')
    document.add_paragraph('lộ trình năm học 2023',style='List Bullet')
    document.add_paragraph('trao thưởng cho học sinh có thành tích suất sắc',style='List Bullet')
    document.add_paragraph('giải thích các khoản thu ',style='List Bullet')
    document.add_heading('các khoản thu:',level=1)
    document.add_paragraph('học phí học tập 2 kì : 20.000.000 VND',style='List Bullet')
    document.add_paragraph('chi phí phát triển tài năng : 5.000.000 VND',style='List Bullet')
    document.add_paragraph('phí phát triển cơ sở vật chất : 1.000.000 VND',style='List Bullet') 
    document.add_heading('\t\t\t Hà Nội ngày 10 tháng 12 năm 2021')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.add_paragraph('\n')
    document.save("name.docx")
    row = file.readline()